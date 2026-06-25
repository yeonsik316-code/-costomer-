"""
수주 잔고 기반 자재 재고 분석 및 공급 리스크 리포트 생성 프로그램

엑셀 4개(수주, BOM, 재고, 자재마스터)를 읽어
제품별/거래처별/납기일정별 수주 잔고를 분석하고,
납기 시 필요 자재 재고를 산출한 뒤 부족분과 월별 발주 비용을 계산하여
OpenAI API로 공급 리스크 분석 워드 리포트를 생성합니다.
"""

import json
import os
import sys
from datetime import datetime
from pathlib import Path

import pandas as pd
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# ---------------------------------------------------------------------------
# 설정
# ---------------------------------------------------------------------------
BASE_DIR = Path(__file__).parent
ORDER_FILE = BASE_DIR / "order_data.xlsx"
BOM_FILE = BASE_DIR / "bom_data.xlsx"
INVENTORY_FILE = BASE_DIR / "inventory_data.xlsx"
MATERIAL_MASTER_FILE = BASE_DIR / "material_master.xlsx"
OUTPUT_DIR = BASE_DIR / "output"
OPENAI_MODEL = "gpt-4o-mini"


# ---------------------------------------------------------------------------
# 데이터 로드
# ---------------------------------------------------------------------------
def load_data() -> dict[str, pd.DataFrame]:
    """엑셀 4개 파일을 읽어 DataFrame 딕셔너리로 반환합니다."""
    orders = pd.read_excel(ORDER_FILE)
    bom = pd.read_excel(BOM_FILE)
    inventory = pd.read_excel(INVENTORY_FILE)
    material_master = pd.read_excel(MATERIAL_MASTER_FILE)

    orders["수주일"] = pd.to_datetime(orders["수주일"])
    orders["납기요청일"] = pd.to_datetime(orders["납기요청일"])

    return {
        "orders": orders,
        "bom": bom,
        "inventory": inventory,
        "material_master": material_master,
    }


# ---------------------------------------------------------------------------
# 수주 잔고 분석
# ---------------------------------------------------------------------------
def analyze_order_balance(orders: pd.DataFrame) -> dict[str, pd.DataFrame]:
    """수주 잔고를 제품별, 거래처별, 납기일정별로 집계합니다."""
    active = orders[orders["수주상태"] == "진행중"].copy()
    active["납기월"] = active["납기요청일"].dt.to_period("M").astype(str)

    by_product = (
        active.groupby(["제품코드", "제품명", "제품군"], as_index=False)
        .agg(수주건수=("수주번호", "count"), 수주수량=("수주수량", "sum"), 수주금액=("수주금액(원)", "sum"))
        .sort_values("수주금액", ascending=False)
    )

    by_customer = (
        active.groupby(["거래처명", "담당영업"], as_index=False)
        .agg(수주건수=("수주번호", "count"), 수주수량=("수주수량", "sum"), 수주금액=("수주금액(원)", "sum"))
        .sort_values("수주금액", ascending=False)
    )

    by_delivery = (
        active.groupby(["납기요청일", "납기월"], as_index=False)
        .agg(수주건수=("수주번호", "count"), 수주수량=("수주수량", "sum"), 수주금액=("수주금액(원)", "sum"))
        .sort_values("납기요청일")
    )

    return {
        "by_product": by_product,
        "by_customer": by_customer,
        "by_delivery": by_delivery,
        "active_orders": active,
    }


# ---------------------------------------------------------------------------
# 자재 소요량 산출 및 재고 부족 분석
# ---------------------------------------------------------------------------
def calculate_material_requirements(
    active_orders: pd.DataFrame,
    bom: pd.DataFrame,
    inventory: pd.DataFrame,
    material_master: pd.DataFrame,
) -> dict:
    """납기일 기준 자재 소요량을 산출하고 재고 부족분을 계산합니다."""
    # 수주 × BOM 조인 → 자재별 소요량
    order_bom = active_orders.merge(bom, on=["제품코드", "제품명"], how="inner")
    order_bom["자재소요량"] = order_bom["수주수량"] * order_bom["소요수량"]
    order_bom["납기월"] = order_bom["납기요청일"].dt.to_period("M").astype(str)

    # 납기월별 자재 소요량
    monthly_demand = (
        order_bom.groupby(["납기월", "자재코드", "자재명"], as_index=False)["자재소요량"]
        .sum()
        .sort_values(["납기월", "자재코드"])
    )

    # 전체 자재 총소요량
    total_demand = (
        order_bom.groupby(["자재코드", "자재명"], as_index=False)["자재소요량"]
        .sum()
        .rename(columns={"자재소요량": "총소요량"})
    )

    # 현재 재고와 병합
    stock = inventory[["자재코드", "자재명", "현재고수량", "안전재고수량", "재고상태"]].copy()
    demand_stock = total_demand.merge(stock, on=["자재코드", "자재명"], how="left")
    demand_stock["현재고수량"] = demand_stock["현재고수량"].fillna(0)
    demand_stock["안전재고수량"] = demand_stock["안전재고수량"].fillna(0)
    demand_stock["부족수량"] = (demand_stock["총소요량"] - demand_stock["현재고수량"]).clip(lower=0)
    demand_stock["안전재고미달"] = (
        demand_stock["현재고수량"] - demand_stock["안전재고수량"] < 0
    )

    shortage = demand_stock[demand_stock["부족수량"] > 0].copy()

    # 납기일 순 누적 재고 소진 시뮬레이션 (월별 부족 발생 시점)
    running_inventory = dict(
        zip(inventory["자재코드"], inventory["현재고수량"])
    )
    monthly_shortages: list[dict] = []

    sorted_orders = order_bom.sort_values("납기요청일")
    for _, row in sorted_orders.iterrows():
        mat_code = row["자재코드"]
        need = row["자재소요량"]
        available = running_inventory.get(mat_code, 0)

        if need > available:
            deficit = need - available
            monthly_shortages.append(
                {
                    "납기월": row["납기월"],
                    "납기요청일": row["납기요청일"].strftime("%Y-%m-%d"),
                    "수주번호": row["수주번호"],
                    "제품코드": row["제품코드"],
                    "자재코드": mat_code,
                    "자재명": row["자재명"],
                    "부족수량": deficit,
                }
            )
            running_inventory[mat_code] = 0
        else:
            running_inventory[mat_code] = available - need

    monthly_shortage_df = pd.DataFrame(monthly_shortages)

    return {
        "order_bom": order_bom,
        "monthly_demand": monthly_demand,
        "total_demand": demand_stock,
        "shortage": shortage,
        "monthly_shortage": monthly_shortage_df,
    }


# ---------------------------------------------------------------------------
# 월별 발주 비용 산출
# ---------------------------------------------------------------------------
def calculate_monthly_procurement_cost(
    monthly_shortage: pd.DataFrame,
    material_master: pd.DataFrame,
) -> pd.DataFrame:
    """월별 자재 부족분에 대한 발주 비용을 산출합니다."""
    if monthly_shortage.empty:
        return pd.DataFrame(
            columns=[
                "납기월", "자재코드", "자재명", "부족수량", "단가(원)",
                "발주수량", "발주비용(원)", "협력사명", "리드타임(일)", "장납기여부",
            ]
        )

    master = material_master[
        ["자재코드", "자재명", "협력사명", "리드타임(일)", "장납기여부", "단가(원)", "최소발주수량"]
    ].copy()

    cost_detail = monthly_shortage.merge(master, on=["자재코드", "자재명"], how="left")

    # 최소발주수량(MOQ) 반영
    cost_detail["발주수량"] = cost_detail.apply(
        lambda r: max(r["부족수량"], r["최소발주수량"]) if pd.notna(r["최소발주수량"]) else r["부족수량"],
        axis=1,
    )
    cost_detail["발주비용(원)"] = cost_detail["발주수량"] * cost_detail["단가(원)"]

    monthly_summary = (
        cost_detail.groupby("납기월", as_index=False)
        .agg(
            부족자재건수=("자재코드", "nunique"),
            총발주수량=("발주수량", "sum"),
            총발주비용=("발주비용(원)", "sum"),
        )
        .sort_values("납기월")
    )

    return cost_detail, monthly_summary


# ---------------------------------------------------------------------------
# 분석 결과를 OpenAI용 텍스트로 변환
# ---------------------------------------------------------------------------
def build_analysis_payload(
    order_balance: dict,
    material_result: dict,
    cost_detail: pd.DataFrame,
    monthly_cost: pd.DataFrame,
) -> str:
    """OpenAI API에 전달할 분석 데이터 JSON 문자열을 생성합니다."""
    def df_to_records(df: pd.DataFrame) -> list:
        records = df.copy()
        for col in records.select_dtypes(include=["datetime64"]).columns:
            records[col] = records[col].astype(str)
        return records.to_dict(orient="records")

    payload = {
        "분석일시": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "수주_제품별": df_to_records(order_balance["by_product"]),
        "수주_거래처별": df_to_records(order_balance["by_customer"]),
        "수주_납기일정별": df_to_records(order_balance["by_delivery"]),
        "자재_총소요_및_재고": df_to_records(material_result["total_demand"]),
        "자재_부족_목록": df_to_records(material_result["shortage"]),
        "월별_자재소요량": df_to_records(material_result["monthly_demand"]),
        "월별_발주비용_상세": df_to_records(cost_detail),
        "월별_발주비용_합계": df_to_records(monthly_cost),
    }
    return json.dumps(payload, ensure_ascii=False, indent=2)


# ---------------------------------------------------------------------------
# OpenAI API로 리포트 본문 생성
# ---------------------------------------------------------------------------
def generate_report_content(client: OpenAI, analysis_data: str) -> str:
    """OpenAI API를 호출하여 공급 리스크 분석 리포트 본문을 생성합니다."""
    system_prompt = """당신은 제조업 공급망 분석 전문가입니다.
제공된 수주·자재·재고·발주비용 데이터를 바탕으로
'공급 리스크 분석 리포트'를 한국어로 작성하세요.

리포트는 다음 섹션을 반드시 포함해야 합니다:
1. Executive Summary (경영진 요약)
2. 수주 현황 분석 (제품별, 거래처별, 납기일정별)
3. 자재 소요 및 재고 현황
4. 재고 부족 자재 분석 (리스크 수준: 높음/중간/낮음)
5. 월별 발주 비용 전망
6. 공급 리스크 요인 (장납기 자재, 해외 협력사, 안전재고 미달 등)
7. 권고 조치사항 (우선순위별)

작성 지침:
- 구체적인 수치와 자재명, 거래처명을 인용하세요.
- 장납기여부=Y 자재와 해외 협력사 자재에 특별히 주의하세요.
- 실무자가 바로 실행할 수 있는 구체적 권고를 제시하세요.
- 마크다운 형식으로 작성하되, ## 로 섹션 제목을, ### 로 하위 제목을 표시하세요.
- 표가 필요한 경우 마크다운 테이블을 사용하세요."""

    response = client.chat.completions.create(
        model=OPENAI_MODEL,
        messages=[
            {"role": "system", "content": system_prompt},
            {
                "role": "user",
                "content": f"다음 데이터를 분석하여 공급 리스크 분석 리포트를 작성해 주세요:\n\n{analysis_data}",
            },
        ],
        temperature=0.3,
        max_tokens=4096,
    )
    return response.choices[0].message.content


# ---------------------------------------------------------------------------
# 워드 문서 생성
# ---------------------------------------------------------------------------
def create_word_report(
    report_content: str,
    order_balance: dict,
    material_result: dict,
    monthly_cost: pd.DataFrame,
    output_path: Path,
) -> None:
    """분석 결과와 AI 생성 본문을 포함한 워드 리포트를 생성합니다."""
    doc = Document()

    # 제목
    title = doc.add_heading("공급 리스크 분석 리포트", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(f"생성일시: {datetime.now().strftime('%Y년 %m월 %d일 %H:%M')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # AI 생성 본문 파싱 (마크다운 → 워드)
    _render_markdown_content(doc, report_content)

    # 부록: 핵심 데이터 테이블
    doc.add_page_break()
    doc.add_heading("부록 A. 수주 잔고 — 제품별", level=1)
    _add_dataframe_table(doc, order_balance["by_product"])

    doc.add_heading("부록 B. 수주 잔고 — 거래처별", level=1)
    _add_dataframe_table(doc, order_balance["by_customer"])

    doc.add_heading("부록 C. 수주 잔고 — 납기일정별", level=1)
    _add_dataframe_table(doc, order_balance["by_delivery"])

    doc.add_heading("부록 D. 자재 부족 목록", level=1)
    shortage_cols = ["자재코드", "자재명", "총소요량", "현재고수량", "부족수량"]
    _add_dataframe_table(doc, material_result["shortage"][shortage_cols])

    doc.add_heading("부록 E. 월별 발주 비용 합계", level=1)
    _add_dataframe_table(doc, monthly_cost)

    doc.save(output_path)


def _render_markdown_content(doc: Document, content: str) -> None:
    """간단한 마크다운 텍스트를 워드 문서 요소로 변환합니다."""
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            _add_markdown_table(doc, table_lines)
            continue
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line[0].isdigit() and ". " in line[:4]:
            doc.add_paragraph(line.split(". ", 1)[1], style="List Number")
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*"))
            run.bold = True
        else:
            doc.add_paragraph(line)

        i += 1


def _add_markdown_table(doc: Document, table_lines: list[str]) -> None:
    """마크다운 테이블을 워드 테이블로 변환합니다."""
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        if all(set(c) <= {"-", ":", " "} for c in cells):
            continue
        rows.append(cells)

    if not rows:
        return

    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < col_count:
                table.rows[r_idx].cells[c_idx].text = cell_text


def _add_dataframe_table(doc: Document, df: pd.DataFrame) -> None:
    """DataFrame을 워드 테이블로 추가합니다."""
    if df.empty:
        doc.add_paragraph("(데이터 없음)")
        return

    display_df = df.copy()
    for col in display_df.columns:
        if display_df[col].dtype in ["float64", "int64"]:
            if "금액" in col or "비용" in col:
                display_df[col] = display_df[col].apply(lambda x: f"{x:,.0f}")
            else:
                display_df[col] = display_df[col].apply(lambda x: f"{x:,.0f}" if pd.notna(x) else "")

    table = doc.add_table(rows=1 + len(display_df), cols=len(display_df.columns))
    table.style = "Table Grid"

    for c_idx, col_name in enumerate(display_df.columns):
        cell = table.rows[0].cells[c_idx]
        cell.text = str(col_name)
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for r_idx, row in enumerate(display_df.itertuples(index=False)):
        for c_idx, value in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(value)


# ---------------------------------------------------------------------------
# 콘솔 요약 출력
# ---------------------------------------------------------------------------
def print_summary(
    order_balance: dict,
    material_result: dict,
    monthly_cost: pd.DataFrame,
) -> None:
    """분석 결과 요약을 콘솔에 출력합니다."""
    sys.stdout.reconfigure(encoding="utf-8")

    print("\n" + "=" * 70)
    print("  공급 리스크 분석 — 실행 결과 요약")
    print("=" * 70)

    active = order_balance["active_orders"]
    print(f"\n[수주 현황] 진행중 수주 {len(active)}건 / 총 수량 {active['수주수량'].sum():,} / "
          f"총 금액 {active['수주금액(원)'].sum():,.0f}원")

    print("\n[제품별 수주 Top 5]")
    for _, row in order_balance["by_product"].head(5).iterrows():
        print(f"  {row['제품명']} ({row['제품코드']}): {row['수주수량']:,}대 / {row['수주금액']:,.0f}원")

    print("\n[거래처별 수주 Top 5]")
    for _, row in order_balance["by_customer"].head(5).iterrows():
        print(f"  {row['거래처명']}: {row['수주건수']}건 / {row['수주금액']:,.0f}원")

    shortage = material_result["shortage"]
    print(f"\n[자재 부족] {len(shortage)}개 자재 부족")
    for _, row in shortage.head(10).iterrows():
        print(f"  {row['자재명']} ({row['자재코드']}): "
              f"소요 {row['총소요량']:,.0f} / 재고 {row['현재고수량']:,.0f} / "
              f"부족 {row['부족수량']:,.0f}")

    print("\n[월별 발주 비용]")
    for _, row in monthly_cost.iterrows():
        print(f"  {row['납기월']}: {row['부족자재건수']}건 / {row['총발주비용']:,.0f}원")

    total_cost = monthly_cost["총발주비용"].sum() if not monthly_cost.empty else 0
    print(f"\n  ▶ 총 예상 발주비용: {total_cost:,.0f}원")
    print("=" * 70 + "\n")


# ---------------------------------------------------------------------------
# 메인 실행
# ---------------------------------------------------------------------------
def main() -> None:
    load_dotenv(BASE_DIR / ".env")
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        print("오류: .env 파일에 OPENAI_API_KEY를 설정해 주세요.")
        sys.exit(1)

    print("데이터 로드 중...")
    data = load_data()

    print("수주 잔고 분석 중...")
    order_balance = analyze_order_balance(data["orders"])

    print("자재 소요량 및 재고 부족 분석 중...")
    material_result = calculate_material_requirements(
        order_balance["active_orders"],
        data["bom"],
        data["inventory"],
        data["material_master"],
    )

    print("월별 발주 비용 산출 중...")
    cost_detail, monthly_cost = calculate_monthly_procurement_cost(
        material_result["monthly_shortage"],
        data["material_master"],
    )

    print_summary(order_balance, material_result, monthly_cost)

    print("OpenAI API로 리포트 생성 중...")
    client = OpenAI(api_key=api_key)
    analysis_payload = build_analysis_payload(
        order_balance, material_result, cost_detail, monthly_cost
    )
    report_content = generate_report_content(client, analysis_payload)

    OUTPUT_DIR.mkdir(exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = OUTPUT_DIR / f"공급리스크분석리포트_{timestamp}.docx"

    print("워드 문서 생성 중...")
    create_word_report(
        report_content,
        order_balance,
        material_result,
        monthly_cost,
        output_path,
    )

    # 분석 데이터 JSON도 저장
    json_path = OUTPUT_DIR / f"분석데이터_{timestamp}.json"
    json_path.write_text(analysis_payload, encoding="utf-8")

    print(f"\n완료!")
    print(f"  리포트: {output_path}")
    print(f"  분석데이터: {json_path}")


if __name__ == "__main__":
    main()
